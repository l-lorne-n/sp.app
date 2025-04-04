import streamlit as st
from openai import OpenAI
from fpdf import FPDF
from docx import Document
import datetime
import re
import unicodedata
import base64
import os

# ✅ 配置 DeepSeek 接口
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),  # ← 换成你自己的deepseek api #api_key="OPENAI_API_KEY"（本地自己简单配置版）
    base_url="https://api.deepseek.com/v1"
)

# === 读取病例文本 ===
def load_case(file):
    return file.read().decode("utf-8")

# === 构建“病人角色”的提示词 ===
def build_system_prompt(case_text, persona="平和"):
    return f"""
你是一位“标准化病人”，请根据以下病例模拟与医学生的对话。

你的性格特点是：【{persona}】。请在回答中体现这种性格特征。

请遵循以下规则：


- 用第一人称回答，不使用专业术语；
- 仅在被提问时作答，不剧透；
- 回答中体现焦虑、疼痛等自然情绪；
- 如果学生提到“口腔检查”，你只需回复“大夫请看”并等待学生观察图像判断。
- 若学生做出诊断，请根据准确性评分：
    ✅ 完全一致（得分：100）
    ⚠️ 缺失部分要点（得分：80）
    ❌ 明显错误（得分：50 以下）

=== 病例内容 ===
{case_text}
=== 你将扮演病人，现在等待提问 ===
"""

# === 构建评分助手的提示词 ===
def build_scoring_prompt(case_text, student_diagnosis):
    return f"""
你是一位医学教学评分助手，请根据以下病例资料，对医学生给出的诊断内容进行分析和打分。

=== 病例内容 ===
{case_text}

=== 学生的诊断内容 ===
{student_diagnosis}

请按以下标准评分：
- ✅ 完全一致，诊断和治疗都准确、完整（得分：100）
- ⚠️ 缺少次要要点，如医嘱、卫生指导等（得分：80）
- ⚠️ 有关键错误，如治疗方式错误（得分：50）
- ❌ 严重偏差，方向错误或信息缺失严重（得分：30）

请直接输出打分和一句简短的评语。
"""

# === 页面部分 ===
st.set_page_config(page_title="SP 模拟器", page_icon="🧠")
st.title("🧠 SP 模拟器 - 标准化病人训练")
persona = st.selectbox(
    "🧠 选择当前病人的性格特点",
    ["平和", "焦虑", "愤怒", "难过", "害怕", "难以沟通", "小孩"],
    index=0,
    help="用于模拟不同类型的标准化病人"
)
st.session_state.persona = persona
uploaded_file = st.file_uploader("上传病例 txt 文件", type="txt")

uploaded_images = st.file_uploader("上传辅助图片（口腔照/X 光片等）", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="uploaded_images")
# 存入状态（只存一次）
if uploaded_images and "image_files" not in st.session_state:
    st.session_state.image_files = uploaded_images

# 控制是否显示
if "show_images" not in st.session_state:
    st.session_state.show_images = False
# 初始化聊天记录
if "messages" not in st.session_state:
    st.session_state.messages = []





# 如果上传了病例
if uploaded_file:
    case_text = load_case(uploaded_file)
    persona = st.session_state.get("persona", "平和")
    system_prompt = build_system_prompt(case_text, persona)

    if not any(m["role"] == "system" for m in st.session_state.messages):
        st.session_state.messages.insert(0, {
            "role": "system",
            "content": system_prompt
        })
    # 展示对话记录
    for msg in st.session_state.messages[1:]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])



# 医学生输入提问
user_input = st.chat_input("你想问病人什么？")
if user_input:
    if "口腔检查" in user_input:
        st.session_state.show_images = True

    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.markdown(user_input)

    response = client.chat.completions.create(
    model="deepseek-chat",
    messages=st.session_state.messages,
    temperature=0.6
)
    reply = response.choices[0].message.content
    st.session_state.messages.append({"role": "assistant", "content": reply})

    with st.chat_message("assistant"):
        st.markdown(reply)

# ✅ 图片显示逻辑紧随回复之后
if (
    st.session_state.get("show_images") and
    st.session_state.get("image_files") and
    len(st.session_state.image_files) > 0
):
    st.subheader("🖼️ 口腔检查结果 / 辅助影像")
    for img in st.session_state.image_files:
        st.image(img, caption=img.name, use_container_width=True)
    st.divider()
    st.subheader("🦷 口腔图像分析评分")

    image_judgement_input = st.text_area(
        "请根据图像描述你的观察结果（例如：缺失牙位、邻牙倾斜、对颌伸长等）",
        height=120,
        placeholder="例：右下46缺失，邻牙无明显倾斜，牙槽嵴平整..."
    )

    if st.button("🧠 提交图像分析评分"):
        if not image_judgement_input.strip():
            st.warning("⚠️ 请填写你的图像观察结果")
        else:
            image_score_prompt = f"""
    你是一名医学教学评分助手，请根据病例内容对学生的“口腔图像观察判断”进行评分。请不要推理，只对比是否与病例一致。

    【评分标准】：
    ✅ 完全一致：得分 100  
    ⚠️ 部分遗漏或有小误：得分 80  
    ❌ 明显错误或严重缺失：得分 50 或以下

    【病例内容】：
    {case_text}

    【学生的图像判断描述】：
    {image_judgement_input}

    请按照下列格式作答：
    ✅ 图像判断评分（100分）：简要点评
    """

            image_response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "system", "content": image_score_prompt}],
                temperature=0.4
            )
            st.session_state.image_score = image_response.choices[0].message.content

    if st.session_state.get("image_score"):
        st.success("✅ 图像评分结果：")
        st.markdown(st.session_state.image_score)


with st.sidebar:
    st.header("📋 教学评分区")

    # 初始化保存评分结果
    if "diagnosis_score" not in st.session_state:
        st.session_state.diagnosis_score = ""
    if "conversation_score" not in st.session_state:
        st.session_state.conversation_score = ""

    st.subheader("🩺 分项诊断评分")

    diagnosis_input = st.text_area("🧾 初步诊断", height=100)
    plan_input = st.text_area("📋 治疗计划", height=100)
    procedure_input = st.text_area("🔧 治疗过程", height=100)

    if st.button("🔍 提交诊断评分"):
        if not any([diagnosis_input.strip(), plan_input.strip(), procedure_input.strip()]):
            st.warning("⚠️ 请填写至少一个部分")
        else:
            combined_input = f"""
你是一名医学教学评分助手，你的任务是根据“原始病例内容”严格打分。  
请不要根据你的推理或外部知识打分，只对比学生的回答和病例是否一致。

【评分规则（必须严格执行）】：
- ✅ 完全一致：学生的回答与病例内容完全相符，无缺失、无错误（得分：100）
- ⚠️ 部分匹配：若存在轻微遗漏，如缺少部分术语/步骤/指导建议（得分：80）
- ❌ 关键错误：若诊断或治疗方式和病例不符，或缺失关键要素（得分：50 或以下）

【原始病例内容】：
{case_text}

【学生提交内容】：
1. 初步诊断：
{diagnosis_input}

2. 治疗计划：
{plan_input}

3. 治疗过程：
{procedure_input}

请你严格对比三项内容，并以如下格式输出评分：
✅ 初步诊断（100分）：简要评语  
⚠️ 治疗计划（80分）：简要评语  
❌ 治疗过程（30分）：简要评语
"""
            score_response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "system", "content": combined_input}],
                temperature=0.5
            )
            st.session_state.diagnosis_score = score_response.choices[0].message.content

    if st.session_state.diagnosis_score:
        st.markdown("✅ **诊断评分结果：**")
        st.markdown(st.session_state.diagnosis_score)

    st.divider()
    st.subheader("💬 沟通整体评分")

    if st.button("🧾 对话整体评分"):
        conversation = ""
        for msg in st.session_state.messages[1:]:
            role = "医学生" if msg["role"] == "user" else "病人"
            conversation += f"{role}：{msg['content']}\n"

        final_prompt = f"""
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与患者交谈过程中的表现。

请关注以下方面：
- 语言是否礼貌、有尊重；
- 是否体现关心患者、缓解焦虑；
- 是否循序渐进、提问有逻辑；
- 没有打断、不操之过急；
- 能否体现出对患者情绪和需求的理解。

请给出：
1. 整体评价（文字简评）；
2. 总体得分（100 分为满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
"""

        final_response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "system", "content": final_prompt}],
            temperature=0.5
        )
        st.session_state.conversation_score = final_response.choices[0].message.content

    if st.session_state.conversation_score:
        st.markdown("✅ **沟通评分结果：**")
        st.markdown(st.session_state.conversation_score)

def clean_text(text):
    if not isinstance(text, str):
        return ""
    return "".join(
        ch for ch in text
        if unicodedata.category(ch)[0] != "C" and ord(ch) < 65536
    ).strip()

def generate_word_report(messages, diagnosis_score, conversation_score, image_score):
    from docx import Document
    import datetime
    import os

    # 创建本地可写目录
    save_dir = "Word_report"
    os.makedirs(save_dir, exist_ok=True)

    doc = Document()
    doc.add_heading("标准化病人训练报告", 0)

    doc.add_heading("对话记录", level=1)
    for msg in messages[1:]:
        role = "医学生" if msg["role"] == "user" else "病人"
        doc.add_paragraph(f"{role}：{msg['content']}")

    doc.add_heading("分项评分", level=1)
    doc.add_paragraph(f"🩺 初步诊断评分：{diagnosis_score}")
    doc.add_paragraph(f"💬 沟通表现评分：{conversation_score}")
    doc.add_paragraph(f"🖼 图像判断评分：{image_score}")

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(save_dir, f"SP训练报告_{timestamp}.docx")

    doc.save(file_path)
    return file_path

if st.button("📄 一键生成 Word 报告"):
    docx_path = generate_word_report(
        messages=st.session_state.messages,
        diagnosis_score=st.session_state.get("diagnosis_score", "无"),
        conversation_score=st.session_state.get("conversation_score", "无"),
        image_score=st.session_state.get("image_score", "无"),
    )
    with open(docx_path, "rb") as f:
        st.download_button("📥 下载 Word 报告", f, file_name="SP训练报告.docx")