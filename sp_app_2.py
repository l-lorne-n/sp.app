import streamlit as st
from openai import OpenAI
from fpdf import FPDF
from docx import Document
import datetime
import re
import unicodedata
import base64
import os
import re, unicodedata

# 会匹配：诊断: 诊断：Diagnosis: diagnosis：
PREFIX_RE = re.compile(r"^\s*(诊断|diagnosis)\s*[:：]\s*", flags=re.IGNORECASE)

def normalize(text: str) -> str:
    """
    1. 去掉 '诊断:' 'Diagnosis:' 等前缀
    2. Unicode 归一化，消除全/半角差异
    3. 去除所有空白并转为小写
    """
    if not isinstance(text, str):
        return ""
    text = PREFIX_RE.sub("", text)          # 去前缀
    text = unicodedata.normalize("NFKC", text)
    return re.sub(r"\s+", "", text.lower())
def generate_word_report(messages, diagnosis_score, conversation_score):
    from docx import Document
    import datetime
    import os

    # 创建本地可写目录
    save_dir = "Word_report"
    os.makedirs(save_dir, exist_ok=True)

    doc = Document()
    doc.add_heading("标准化病人训练报告", 0)

    doc.add_heading("对话记录", level=1)
    for msg in messages[1:]:
        role = "医学生" if msg["role"] == "user" else "病人"
        doc.add_paragraph(f"{role}：{msg['content']}")

    doc.add_heading("分项评分", level=1)
    doc.add_paragraph(f"🩺 初步诊断评分：{diagnosis_score}")
    doc.add_paragraph(f"💬 沟通表现评分：{conversation_score}")
    #doc.add_paragraph(f"🖼 图像判断评分：{image_score}")

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(save_dir, f"SP训练报告_{timestamp}.docx")

    doc.save(file_path)
    return file_path

# ✅ 配置 DeepSeek 接口
client = OpenAI(
    api_key="sk-fd7dfdc16ab5447e8e280bdde625213f",  # ← 换成你自己的
    base_url="https://api.deepseek.com/v1"
)
# === 不同病种打分模板 === 包含：牙痛，牙松动，牙龈肥大，口腔黏膜白色斑块，颌面部肿痛，牙龈出血
SCORING_PROMPT_TEMPLATES = {
    "牙痛": """
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与牙痛患者交谈过程中的表现。

请重点关注以下内容是否被涵盖：

（一）现病史：
1. 是否询问牙痛的性质和程度；
2. 疼痛的出现时间和持续时间；
3. 疼痛部位及是否有放射；
4. 是否询问诱发、加重或缓解因素；
5. 是否了解过去的治疗经历和效果；

（二）既往史：
6. 是否询问患牙的治疗、修复、牙外伤和正畸史；

（三）全身情况：
7. 是否询问系统性疾病（如心脏病、高血压、糖尿病）；
8. 是否考虑女性特殊情况（如经期）；

请给出以下内容：
1. 整体评价（简要文字说明）；
2. 总体得分（100 分满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
""",
    "牙松动":"""    
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与牙松动患者交谈过程中的表现。

请重点关注以下内容是否被涵盖：

（一）现病史：
1. 是否询问牙齿松动的起病时间和诱因；
2. 松动的部位，是单颗还是多颗牙；
3. 是否伴有牙龈出血（时间/场景）、红肿等症状；
4. 是否询问是否有咬合痛、外伤史、咬物时硌伤；
5. 是否关注局部肿物或颌骨膨隆，及其出现时间；

（二）既往史：
6. 是否有牙周治疗史、正畸史；
7. 是否了解颌面部疾病或邻近器官病史；
8. 是否询问过往是否有牙因松动脱落；
9. 是否判断患者是否处于替牙期（如儿童）；

（三）全身情况：
10. 是否询问是否患有糖尿病、冠心病、高血压、肿瘤等系统疾病；
11. 是否了解家族中是否有牙齿早失史；

请给出以下内容：
1. 整体评价（简要文字说明）；
2. 总体得分（100 分为满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
""",
    "牙龈肥大": """
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与牙龈肥大患者交谈过程中的表现。

请重点关注以下内容是否被涵盖：

（一）现病史：
1. 是否明确牙龈肥大的部位、范围和持续时间；
2. 是否询问癫痫、高血压、肾移植史，以及服药史；
3. 是否了解肥大牙龈是否影响咀嚼；
4. 是否询问牙龈是否易出血（是否自发、刷牙时、是否不易止住）；
5. 是否关注牙龈乳头是否瘤样肥大，并进一步了解大小、范围及是否妊娠相关；
6. 是否询问牙齿萌出困难史或家族遗传史；

（二）既往史：
7. 是否有牙龈出血史；
8. 是否了解有无白血病病史；
9. 是否询问口腔颌面部其他相关疾病史；

请给出以下内容：
1. 整体评价（简要文字说明）；
2. 总体得分（100 分为满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
""",
    "口腔黏膜白色斑块": """
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与口腔黏膜白色斑块患者交谈过程中的表现。

请重点关注以下内容是否被涵盖：

（一）现病史：
1. 是否询问外伤史；
2. 是否了解吸烟、饮酒及嚼槟榔习惯；
3. 是否关注白色念珠菌感染情况（如女性是否有阴道炎史）；
4. 是否询问日晒情况（如是否长期户外工作，如农民）；
5. 是否明确白色病损的发病时间、部位、性质（是否对称、粗糙或光滑）、既往治疗史（如是否使用免疫抑制剂）；

（二）既往史：
6. 是否了解 HIV 感染史、精神创伤史、梅毒病史；

（三）全身情况：
7. 是否关注皮肤病、HIV 相关其他疾病；
8. 是否询问家族遗传史；
9. 是否注意使用亲切语气保护患者隐私，尤其在涉及性传播疾病时注意表达方式及医生自身安全；

请给出以下内容：
1. 整体评价（简要文字说明）；
2. 总体得分（100 分为满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
""",
    "颌面部肿痛": """
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与颌面部肿痛患者交谈过程中的表现。

请重点关注以下内容是否被涵盖：

（一）现病史：
1. 是否询问肿痛出现的时间及发展过程；
2. 是否了解肿痛范围是否扩大，有无全身反应；
3. 是否询问疼痛性质（如钝痛、跳痛、压痛）；
4. 是否关注肿物质地（松软/坚硬）、皮肤颜色变化；
5. 是否询问是否伴随功能障碍，如张口受限、吞咽困难等；

（二）既往史：
6. 是否了解肿痛是否反复发作；
7. 是否询问外伤史、手术史、过敏史及其他治疗史；
8. 是否了解以往治疗的效果如何；

（三）全身情况：
9. 是否关注体温变化、发热、寒战、疲倦、无力、食欲不振等全身表现；

请给出以下内容：
1. 整体评价（简要文字说明）；
2. 总体得分（100 分为满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
""",
    "牙龈出血": """
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与牙龈出血患者交谈过程中的表现。

请重点关注以下内容是否被涵盖：

（一）现病史：
1. 是否明确出血的部位、时间（前牙/后牙、单个/多个）；
2. 是否了解是自发出血还是受刺激后出血，是否能自行止血、止血方式；
3. 是否询问出血量；
4. 是否询问是否伴有牙龈疼痛；
5. 是否关注牙龈是否肿胀；
6. 是否询问是否有瘤样物形成；
7. 是否询问是否有牙齿松动或脱落；
8. 是否了解近期是否有紧张、劳累、熬夜等诱因；

（二）既往史：
9. 是否询问以往是否有牙龈出血或身体其他部位出血不易止的情况；
10. 是否了解白血病病史；

（三）全身情况：
11. 是否询问是否处于妊娠、月经期、更年期；
12. 是否有高血压病史；
13. 是否服用抗凝药物（如阿司匹林、华法林等）；

请给出以下内容：
1. 整体评价（简要文字说明）；
2. 总体得分（100 分为满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
""",
    "口腔黏膜溃疡": """
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与“口腔黏膜溃疡”患者交谈过程中的表现。

请重点关注以下内容是否被涵盖：

（一）现病史：
1. 是否询问溃疡是否周期性发作；
2. 溃疡持续的时间；
3. 溃疡的大小、深浅、数目；
4. 是否询问溃疡的好发部位及伴随症状（如发热、淋巴结肿大）；
5. 是否了解溃疡疼痛情况及相关病史；
6. 是否询问是否存在创伤诱因（如残根、残冠、错牙合畸形、不良修复体）；

（二）既往史：
7. 是否询问结核病史、药物过敏史、恶性肿瘤史；

（三）全身情况：
8. 是否关注全身相关症状（如外生殖器溃疡、皮肤病、眼病发作情况）；

请给出以下内容：
1. 整体评价（简要文字说明）；
2. 总体得分（100 分满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
""",






}

# 默认打分模板
default_prompt = """
你是一名医学沟通评价专家，请根据以下对话，评价医学生在与患者交谈过程中的表现。

请关注以下方面：
- 语言是否礼貌、有尊重；
- 是否体现关心患者、缓解焦虑；
- 是否循序渐进、提问有逻辑；
- 没有打断、不操之过急；
- 能否体现出对患者情绪和需求的理解。

请给出：
1. 整体评价（文字简评）；
2. 总体得分（100 分为满分）；
3. 可改进建议（若有）。

=== 对话记录 ===
{conversation}
"""

# === 读取病例文本并提取病种 ===
# === 读取病例文本并提取病种和临床检查提示 ===
def load_case(file):
    text = file.read().decode("utf-8")
    lines = text.strip().splitlines()
    disease_type = lines[0].strip() if len(lines) > 0 else "其他"
    clinical_hint = lines[1].strip() if len(lines) > 1 else ""
    case_text = "\n".join(lines[2:-2]).strip() if len(lines) > 4 else ""
    diagnosis = lines[-2].strip() if len(lines) >= 2 else ""
    plan = lines[-1].strip() if len(lines) >= 1 else ""
    return disease_type, clinical_hint, case_text, diagnosis, plan

#初始化状态变量，避免首次访问报错
for key in ["diagnosis_score", "conversation_score", "case_text", "clinical_hint", "diagnosis_answer", "plan_answer"]:
    if key not in st.session_state:
        st.session_state[key] = ""

# === 构建“病人角色”的提示词 ===
def build_system_prompt(case_text, persona="平和"):
    return f"""
你是一位“标准化病人”，请根据以下病例模拟与医学生的对话。

你的性格特点是：【{persona}】。请在回答中体现这种性格特征。

请遵循以下规则：


- 用第一人称回答，不使用专业术语；
- 仅在被提问时作答，不剧透；
- 回答中体现焦虑、疼痛等自然情绪；
- 如果学生提到“口腔检查”，你只需回复“大夫请看”并等待学生观察图像判断。
- 若学生做出诊断，请根据准确性评分：
    ✅ 完全一致（得分：100）
    ⚠️ 缺失部分要点（得分：80）
    ❌ 明显错误（得分：50 以下）

=== 病例内容 ===
{case_text}
=== 你将扮演病人，现在等待提问 ===
"""

# === 构建评分助手的提示词 ===
def build_scoring_prompt(case_text, student_diagnosis):
    return f"""
你是一位医学教学评分助手，请根据以下病例资料，对医学生给出的诊断内容进行分析和打分。

=== 病例内容 ===
{case_text}

=== 学生的诊断内容 ===
{student_diagnosis}

请按以下标准评分：
- ✅ 完全一致，诊断和治疗都准确、完整（得分：100）
- ⚠️ 缺少次要要点，如医嘱、卫生指导等（得分：80）
- ⚠️ 有关键错误，如治疗方式错误（得分：50）
- ❌ 严重偏差，方向错误或信息缺失严重（得分：30）

请直接输出打分和一句简短的评语。
"""

# === 页面部分 ===
st.set_page_config(page_title="医问灵犀", page_icon="🧠")
st.title("🧠 医问灵犀")
persona = st.selectbox(
    "🧠 选择当前病人的性格特点",
    ["平和", "焦虑", "愤怒", "难过", "害怕", "难以沟通", "小孩"],
    index=0,
    help="用于模拟不同类型的标准化病人"
)
st.session_state.persona = persona
uploaded_file = st.file_uploader("上传病例 txt 文件", type="txt")

uploaded_images = st.file_uploader("上传辅助图片（口腔照/X 光片等）", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="uploaded_images")
# 存入状态（只存一次）
if uploaded_images and "image_files" not in st.session_state:
    st.session_state.image_files = uploaded_images

# 控制是否显示
if "show_images" not in st.session_state:
    st.session_state.show_images = False
# 初始化聊天记录
if "messages" not in st.session_state:
    st.session_state.messages = []





# 如果上传了病例
if uploaded_file:
    # 如果上传了病例
    if uploaded_file is not None:
        # 通过文件名或内容哈希来判断是否是新文件
        if st.session_state.get("last_file_name") != uploaded_file.name:
            disease_type, clinical_hint, case_text, diagnosis, plan_answer = load_case(uploaded_file)

            st.session_state.update({
                "last_file_name": uploaded_file.name,
                "disease_type": disease_type,
                "clinical_hint": clinical_hint,
                "case_text": case_text,
                "diagnosis": diagnosis,
                "plan_answer": plan_answer
            })
            # 3. **重置聊天记录**，并插入新的 system prompt
            st.session_state.messages = []            # 清空旧对话

    # 使用时从 session_state 读取
    case_text = st.session_state.get("case_text", "")
    clinical_hint = st.session_state.get("clinical_hint", "")
    diagnosis = st.session_state.get("diagnosis", "")
    plan_answer = st.session_state.get("plan_answer", "")

    # 显示临床信息（确保 txt 已加载）
    if st.session_state.get("clinical_hint"):
        st.markdown("### 🩻 临床检查信息：")
        st.info(st.session_state["clinical_hint"])
    # 后续使用 case_text 替代原来的 load_case() 返回值
    persona = st.session_state.get("persona", "平和")
    system_prompt = build_system_prompt(case_text, persona)

    if not any(m["role"] == "system" for m in st.session_state.messages):
        st.session_state.messages.insert(0, {
            "role": "system",
            "content": system_prompt
        })
    # 展示对话记录
    for msg in st.session_state.messages[1:]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])



# 医学生输入提问
user_input = st.chat_input("你想问病人什么？")
if user_input:
    if "口腔检查" in user_input:
        st.session_state.show_images = True

    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.markdown(user_input)

    response = client.chat.completions.create(
    model="deepseek-chat",
    messages=st.session_state.messages,
    temperature=0.6
)
    reply = response.choices[0].message.content
    st.session_state.messages.append({"role": "assistant", "content": reply})

    with st.chat_message("assistant"):
        st.markdown(reply)

# ✅ 图片显示逻辑紧随回复之后
if (
    st.session_state.get("show_images") and
    st.session_state.get("image_files") and
    len(st.session_state.image_files) > 0
):
    st.subheader("🖼️ 口腔检查结果 / 辅助影像")
    for img in st.session_state.image_files:
        st.image(img, caption=img.name, use_container_width=True)
    # st.divider()
    # st.subheader("🦷 口腔图像分析评分")

    # image_judgement_input = st.text_area(
    #     "请根据图像描述你的观察结果（例如：缺失牙位、邻牙倾斜、对颌伸长等）",
    #     height=120,
    #     placeholder="例：右下46缺失，邻牙无明显倾斜，牙槽嵴平整..."
    # )

    # if st.button("🧠 提交图像分析评分"):
    #     if not image_judgement_input.strip():
    #         st.warning("⚠️ 请填写你的图像观察结果")
    #     else:
    #         image_score_prompt = f"""
    # 你是一名医学教学评分助手，请根据病例内容对学生的“口腔图像观察判断”进行评分。请不要推理，只对比是否与病例一致。

    # 【评分标准】：
    # ✅ 完全一致：得分 100  
    # ⚠️ 部分遗漏或有小误：得分 80  
    # ❌ 明显错误或严重缺失：得分 50 或以下

    # 【病例内容】：
    # {case_text}

    # 【学生的图像判断描述】：
    # {image_judgement_input}

    # 请按照下列格式作答：
    # ✅ 图像判断评分（100分）：简要点评
    # """

    #         image_response = client.chat.completions.create(
    #             model="deepseek-chat",
    #             messages=[{"role": "system", "content": image_score_prompt}],
    #             temperature=0.4
    #         )
    #         st.session_state.image_score = image_response.choices[0].message.content

    # if st.session_state.get("image_score"):
    #     st.success("✅ 图像评分结果：")
    #     st.markdown(st.session_state.image_score)


# with st.sidebar:
#     st.header("📋 教学评分区")

#     # 初始化保存评分结果
#     if "diagnosis_score" not in st.session_state:
#         st.session_state.diagnosis_score = ""
#     if "conversation_score" not in st.session_state:
#         st.session_state.conversation_score = ""

#     st.subheader("🩺 分项诊断评分")

#     diagnosis_input = st.text_area("🧾 初步诊断", height=100)
#     plan_input = st.text_area("📋 治疗计划", height=100)
#     procedure_input = st.text_area("🔧 治疗过程", height=100)

#     if st.button("🔍 提交诊断评分"):
#         if not any([diagnosis_input.strip(), plan_input.strip(), procedure_input.strip()]):
#             st.warning("⚠️ 请填写至少一个部分")
#         else:
#             combined_input = f"""
# 你是一名医学教学评分助手，你的任务是根据“原始病例内容”严格打分。  
# 请不要根据你的推理或外部知识打分，只对比学生的回答和病例是否一致。

# 【评分规则（必须严格执行）】：
# - ✅ 完全一致：学生的回答与病例内容完全相符，无缺失、无错误（得分：100）
# - ⚠️ 部分匹配：若存在轻微遗漏，如缺少部分术语/步骤/指导建议（得分：80）
# - ❌ 关键错误：若诊断或治疗方式和病例不符，或缺失关键要素（得分：50 或以下）

# 【原始病例内容】：
# {case_text}

# 【学生提交内容】：
# 1. 初步诊断：
# {diagnosis_input}

# 2. 治疗计划：
# {plan_input}

# 3. 治疗过程：
# {procedure_input}

# 请你严格对比三项内容，并以如下格式输出评分：
# ✅ 初步诊断（100分）：简要评语  
# ⚠️ 治疗计划（80分）：简要评语  
# ❌ 治疗过程（30分）：简要评语
# """
#             score_response = client.chat.completions.create(
#                 model="deepseek-chat",
#                 messages=[{"role": "system", "content": combined_input}],
#                 temperature=0.5
#             )
#             st.session_state.diagnosis_score = score_response.choices[0].message.content

#     if st.session_state.diagnosis_score:
#         st.markdown("✅ **诊断评分结果：**")
#         st.markdown(st.session_state.diagnosis_score)
# 诊断评分区
with st.sidebar:
    st.header("📋 教学评分区")
    st.subheader("🩺 分项诊断评分")

    # 学生输入
    diagnosis_input = st.text_area("🧾 初步诊断", height=100)
    plan_input      = st.text_area("📋 治疗计划", height=100)

    # 点击按钮
    if st.button("🔍 提交诊断评分"):
        if not any([diagnosis_input.strip(), plan_input.strip()]):
            st.warning("⚠️ 请填写至少一个部分")
        else:
            # -------- 1️⃣ 诊断语义评分（只比两句话） --------
            std_diag = st.session_state.get("diagnosis", "")

            diag_prompt = f"""
你是一名医学教学评分助手，现在仅比较两句话的**语义一致性**并打分。

【评分标准】
- 完全语义一致（同义表达、格式差异均算一致）→ 100 分
- 基本一致，但缺少/多了次要修饰成分 → 80 分
- 关键诊断方向错误或缺失 → 50 分

【病例标准初步诊断】
{std_diag}

【学生初步诊断】
{diagnosis_input}

请按如下格式输出：
✅ 初步诊断（分数）：一句简评
"""
            diag_resp = client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "system", "content": diag_prompt}],
                temperature=0.2
            ).choices[0].message.content.strip()

            # -------- 2️⃣ 治疗计划建议（不打分） --------
            plan_prompt = f"""
你是一名口腔科教学导师，请针对下面“学生治疗计划”给出 1~2 句改进或肯定建议（不必打分）。

【病例标准治疗计划】
{st.session_state.get('plan_answer', '（病例未提供）')}

【学生治疗计划】
{plan_input}
"""
            plan_resp = client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "system", "content": plan_prompt}],
                temperature=0.2
            ).choices[0].message.content.strip()

            # -------- 3️⃣ 合并并存入 session --------
            st.session_state.diagnosis_score = (
                f"{diag_resp}\n"
                f"💡 治疗计划建议：{plan_resp}"
            )

    # 显示结果
    if st.session_state.get("diagnosis_score"):
        st.markdown("✅ **诊断评分结果：**")
        st.markdown(st.session_state.diagnosis_score)

    # 保证沟通评分 key 存在
    if "conversation_score" not in st.session_state:
        st.session_state.conversation_score = ""




    st.divider()
    st.subheader("💬 沟通整体评分")

    if st.button("🧾 对话整体评分"):
        conversation = ""
        for msg in st.session_state.messages[1:]:
            role = "医学生" if msg["role"] == "user" else "病人"
            conversation += f"{role}：{msg['content']}\n"

        disease_type = st.session_state.get("disease_type", "其他")
        prompt_template = SCORING_PROMPT_TEMPLATES.get(disease_type, default_prompt)
        final_prompt = prompt_template.format(conversation=conversation)

        final_response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "system", "content": final_prompt}],
            temperature=0.5
        )
        st.session_state.conversation_score = final_response.choices[0].message.content

    if st.session_state.conversation_score:
        st.markdown("✅ **沟通评分结果：**")
        st.markdown(st.session_state.conversation_score)
        st.divider()
    st.subheader("📄 生成报告")
    if st.button("一键导出 Word"):
        docx_path = generate_word_report(
            messages=st.session_state.messages,
            diagnosis_score=st.session_state.get("diagnosis_score", "无"),
            conversation_score=st.session_state.get("conversation_score", "无"),
            #image_score=st.session_state.get("image_score", "无"),
        )
        with open(docx_path, "rb") as f:
            st.download_button("📥 下载 Word 报告", f, file_name="SP训练报告.docx")

def clean_text(text):
    if not isinstance(text, str):
        return ""
    return "".join(
        ch for ch in text
        if unicodedata.category(ch)[0] != "C" and ord(ch) < 65536
    ).strip()

def generate_word_report(messages, diagnosis_score, conversation_score):
    from docx import Document
    import datetime
    import os

    # 创建本地可写目录
    save_dir = "Word_report"
    os.makedirs(save_dir, exist_ok=True)

    doc = Document()
    doc.add_heading("标准化病人训练报告", 0)

    doc.add_heading("对话记录", level=1)
    for msg in messages[1:]:
        role = "医学生" if msg["role"] == "user" else "病人"
        doc.add_paragraph(f"{role}：{msg['content']}")

    doc.add_heading("分项评分", level=1)
    doc.add_paragraph(f"🩺 初步诊断评分：{diagnosis_score}")
    doc.add_paragraph(f"💬 沟通表现评分：{conversation_score}")
    #doc.add_paragraph(f"🖼 图像判断评分：{image_score}")

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(save_dir, f"SP训练报告_{timestamp}.docx")

    doc.save(file_path)
    return file_path
# === 一键生成 Word 报告 ===
if st.button("📄 一键生成 Word 报告"):
    docx_path = generate_word_report(
        messages=st.session_state.messages,
        diagnosis_score=st.session_state.get("diagnosis_score", "无"),
        conversation_score=st.session_state.get("conversation_score", "无"),
        #image_score=st.session_state.get("image_score", "无"),
    )
    with open(docx_path, "rb") as f:
        st.download_button("📥 下载 Word 报告", f, file_name="SP训练报告.docx")

# with st.sidebar:
#     st.subheader("💬 沟通整体评分")

#     # 初始化评分结果
#     if "conversation_score" not in st.session_state:
#         st.session_state.conversation_score = ""

#     if st.button("🧾 对话整体评分"):
#         conversation = ""
#         for msg in st.session_state.messages[1:]:
#             role = "医学生" if msg["role"] == "user" else "病人"
#             conversation += f"{role}：{msg['content']}\n"

#         disease_type = st.session_state.get("disease_type", "其他")
#         prompt_template = SCORING_PROMPT_TEMPLATES.get(disease_type, default_prompt)
#         final_prompt = prompt_template.format(conversation=conversation)

#         final_response = client.chat.completions.create(
#             model="deepseek-chat",
#             messages=[{"role": "system", "content": final_prompt}],
#             temperature=0.5
#         )
#         st.session_state.conversation_score = final_response.choices[0].message.content

#     if st.session_state.conversation_score:
#         st.success("✅ 沟通评分结果：")
#         st.markdown(st.session_state.conversation_score)